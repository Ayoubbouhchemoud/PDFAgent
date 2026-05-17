"""
pdf_to_docx.py — PDF to DOCX converter.

Reuses the same pdfplumber-based structural extraction as pdf_to_html.py and
maps each logical block to the closest Word-native element:

  Headings          → Bold Normal paragraphs with bottom border (matches PDF)
  Body text         → Normal paragraphs with Calibri font (wrapped lines merged)
  Bullet lists      → List Bullet style (real Word lists, not manual •)
  Code/syntax boxes → bordered single-cell tables (monospace Consolas)
  Side-by-side boxes→ 2-column borderless table, each cell bordered
  Tabular data      → Word table
  Images            → inline embedded images; if side-by-side with text, layout table
  Page content      → one continuous DOCX (no hard page breaks between pages)
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import pdfplumber
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Import shared parsing helpers from pdf_to_html.py (same directory)
# ---------------------------------------------------------------------------
_HERE = Path(__file__).parent
sys.path.insert(0, str(_HERE))

from pdf_to_html import (
    _COL_DETECT_GAP,
    _LINE_TOL,
    _MERGE_GAP,
    _MIN_TABLE_COLS,
    _assign_col,
    _base_font,
    _content_boxes,
    _css_color,
    _detect_col_starts,
    _group_into_lines,
    _is_bold,
    _is_italic,
    _is_mono,
    _is_white_color,
    _merge_into_runs,
    _remap_pua,
    _resolve_indexed_palette,
    _section_rects,
)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _set_cell_borders(cell, color: str = "000000", sz: int = 4):
    """Apply a thin single border on all four sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _clear_cell_borders(cell):
    """Remove all borders from a table cell (for layout tables)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_table_no_borders(table):
    """Remove all borders from a table (used for layout tables)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _add_para_bottom_border(para, color: str = "000000", sz: int = 6):
    """Add a bottom border to a paragraph (used to simulate heading underlines)."""
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run(para, text: str, bold: bool = False, italic: bool = False,
             fontname: str = "Calibri", size: float = 11.0):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = _map_fontname(fontname)
    run.font.size = Pt(size)
    return run


def _map_fontname(fname: str) -> str:
    """Map a PDF font name (e.g. 'ABCDEE+Calibri,Bold') to a Word font name."""
    base = _base_font(fname)
    _MAP = {
        "calibri":  "Calibri",
        "arial":    "Arial",
        "consolas": "Consolas",
        "courier":  "Courier New",
        "times":    "Times New Roman",
        "helvetica":"Arial",
        "verdana":  "Verdana",
        "tahoma":   "Tahoma",
        "symbol":   "Symbol",
    }
    for key, val in _MAP.items():
        if key in base:
            return val
    return "Calibri"


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def _extract_image_bytes(img: dict) -> tuple[bytes, str] | None:
    """Return (raw_bytes, mime_type) for an image dict, or None on failure."""
    stream = img.get("stream")
    if stream is None:
        return None
    try:
        raw = stream.get_data()
    except Exception:
        return None

    filt = str(stream.attrs.get("Filter", ""))
    if "DCT" in filt or "JPEG" in filt:
        return raw, "image/jpeg"

    try:
        from PIL import Image as PILImage

        px_w = int(stream.attrs.get("Width",  0))
        px_h = int(stream.attrs.get("Height", 0))
        if not px_w or not px_h:
            return None

        cs_raw = stream.attrs.get("ColorSpace", img.get("colorspace", ""))
        cs = str(cs_raw)

        if "Gray" in cs:
            mode = "L"; expected = px_w * px_h
        elif "Indexed" in cs or "indexed" in cs.lower():
            mode = "P"; expected = px_w * px_h
        elif "CMYK" in cs:
            mode = "CMYK"; expected = px_w * px_h * 4
        else:
            mode = "RGB"; expected = px_w * px_h * 3

        if len(raw) < expected:
            return None

        pil = PILImage.frombytes(mode, (px_w, px_h), raw[:expected])
        if mode == "P":
            pal = _resolve_indexed_palette(cs_raw)
            if pal:
                pil.putpalette(pal)
            pil = pil.convert("RGB")
        elif mode == "CMYK":
            pil = pil.convert("RGB")

        buf = io.BytesIO()
        pil.save(buf, format="PNG")
        return buf.getvalue(), "image/png"
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Block detection
# ---------------------------------------------------------------------------

def _is_bullet_word(w: dict) -> bool:
    """Return True if this word is a Symbol-font bullet character."""
    fname = w.get("fontname", "")
    text  = _remap_pua(w["text"], fname)
    return "symbol" in fname.lower() and text in ("•", "▪", "◦", "‣")


def _underline_ys(page) -> set[float]:
    """Collect Y positions of thin horizontal filled-black rects (heading underlines)."""
    pw = float(page.width)
    result: set[float] = set()
    for r in page.rects:
        rw = r["x1"] - r["x0"]
        rh = r["bottom"] - r["top"]
        fc = r.get("non_stroking_color")
        if (r.get("fill") and not _is_white_color(fc) and
                rh < 3.0 and rw > 40 and rw < pw * 0.95):
            result.add(r["top"])
    return result


def _find_box_pairs(boxes: list[dict]) -> tuple[list[tuple], list[dict]]:
    """Split content boxes into side-by-side pairs and standalone singles."""
    pairs: list[tuple] = []
    used:  set[int]    = set()
    for i, b1 in enumerate(boxes):
        if i in used:
            continue
        for j, b2 in enumerate(boxes):
            if j <= i or j in used:
                continue
            y_overlap = min(b1["bottom"], b2["bottom"]) - max(b1["top"], b2["top"])
            if y_overlap > 10 and (b1["x1"] < b2["x0"] or b2["x1"] < b1["x0"]):
                left  = b1 if b1["x0"] < b2["x0"] else b2
                right = b2 if b1["x0"] < b2["x0"] else b1
                pairs.append((left, right))
                used.add(i); used.add(j)
    singles = [boxes[k] for k in range(len(boxes)) if k not in used]
    return pairs, singles


def _line_height(line: list[dict]) -> float:
    """Estimate the line height from font size."""
    if not line:
        return 14.0
    return float(line[0].get("size", 11)) * 1.3


def _group_into_paragraphs(lines: list[list[dict]], pw: float) -> list[list[list[dict]]]:
    """
    Merge consecutive visually-wrapped lines into logical paragraphs.

    A line wraps into the previous one when ALL of:
      - The vertical gap is within normal line-spacing (≤ 1.8 × line-height)
      - The left margins match (within 8pt)
      - The previous line ends near the right margin (gap < 100pt), meaning it
        was forced to wrap — short lines that end early are their own paragraph
      - Neither line is a bullet
    """
    if not lines:
        return []
    groups: list[list[list[dict]]] = [[lines[0]]]
    for curr in lines[1:]:
        prev = groups[-1][-1]
        gap     = curr[0]["top"] - prev[0]["top"]
        prev_x0 = min(w["x0"] for w in prev)
        curr_x0 = min(w["x0"] for w in curr)
        prev_x1 = max(w.get("x1", w["x0"]) for w in prev)
        lh      = _line_height(prev)
        is_wrap = (
            gap <= lh * 1.8 and
            abs(curr_x0 - prev_x0) < 8 and
            (pw - prev_x1) < 100 and          # prev line reached near right margin
            not _is_bullet_word(curr[0]) and
            not _is_bullet_word(prev[0])
        )
        if is_wrap:
            groups[-1].append(curr)
        else:
            groups.append([curr])
    return groups


# ---------------------------------------------------------------------------
# DOCX writers
# ---------------------------------------------------------------------------

def _write_runs(para, runs: list[dict]):
    """Write a list of run-dicts (from _merge_into_runs) into a paragraph."""
    prev_x1: float | None = None
    for run in runs:
        fname = run.get("fontname", "")
        text  = _remap_pua(run["text"], fname)
        # Preserve inter-run word spacing: prepend a space when there's a gap
        if prev_x1 is not None and run.get("x0", prev_x1) - prev_x1 > 0.5:
            text = " " + text
        prev_x1 = run.get("x1", run.get("x0", 0))
        _add_run(
            para, text,
            bold     = _is_bold(fname),
            italic   = _is_italic(fname),
            fontname = fname,
            size     = float(run.get("size", 11)),
        )


def _write_para_lines(para, lines: list[list[dict]]) -> None:
    """Write one or more merged lines into a single paragraph with space joins."""
    for li, line in enumerate(lines):
        if li > 0:
            para.add_run(" ")   # word-space between wrapped lines
        _write_runs(para, _merge_into_runs(line))


def _write_bullet(doc: Document, line: list[dict]) -> None:
    """Write a single bullet line as a List Bullet paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    text_words = [w for w in line if not _is_bullet_word(w)]
    _write_runs(para, _merge_into_runs(text_words))


def _write_code_box(doc: Document, words: list[dict], box: dict) -> None:
    """Write a single code box as a bordered, monospace single-cell table."""
    if not words:
        return
    table = doc.add_table(rows=1, cols=1)
    _set_table_no_borders(table)
    cell = table.cell(0, 0)
    _set_cell_borders(cell)
    # Remove the default empty paragraph inside the cell
    default_para = cell.paragraphs[0]
    default_para._element.getparent().remove(default_para._element)

    for line in _group_into_lines(words):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        _write_runs(p, _merge_into_runs(line))


def _write_side_by_side(doc: Document,
                         left_words:  list[dict], left_box:  dict,
                         right_words: list[dict], right_box: dict,
                         left_label:  str = "",   right_label: str = "") -> None:
    """Write a pair of side-by-side code boxes as a 2-column Word table."""
    # Label row above the code pair
    if left_label or right_label:
        lbl = doc.add_paragraph(style="Normal")
        lbl.add_run(left_label).font.name  = "Calibri"
        lbl.add_run("\t")
        lbl.add_run(right_label).font.name = "Calibri"

    table = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(table)
    row = table.rows[0]

    for cell_idx, (words, box) in enumerate(
            [(left_words, left_box), (right_words, right_box)]):
        cell = row.cells[cell_idx]
        _set_cell_borders(cell)
        default_para = cell.paragraphs[0]
        default_para._element.getparent().remove(default_para._element)
        for line in _group_into_lines(words):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            _write_runs(p, _merge_into_runs(line))


def _write_table(doc: Document, words: list[dict],
                 col_starts: list[float], y_top: float) -> None:
    """Write a detected tabular region as a Word table."""
    lines = _group_into_lines(words)
    if not lines:
        return

    n = len(col_starts)

    from pdf_to_html import _DIGIT_RE

    header_lines: list[list[dict]] = []
    data_lines:   list[list[dict]] = []
    in_header = True
    for line in lines:
        col0 = [w for w in line if _assign_col(w["x0"], col_starts) == 0]
        col0_text = " ".join(w["text"] for w in col0)
        if in_header:
            if col0 and _DIGIT_RE.search(col0_text):
                in_header = False
                data_lines.append(line)
            else:
                header_lines.append(line)
        else:
            data_lines.append(line)

    logical_rows: list[list[list[dict]]] = []
    cur: list[list[dict]] = []
    for line in data_lines:
        has_col0 = any(_assign_col(w["x0"], col_starts) == 0 for w in line)
        if has_col0 and cur:
            logical_rows.append(cur)
            cur = []
        cur.append(line)
    if cur:
        logical_rows.append(cur)

    total_rows = (1 if header_lines else 0) + len(logical_rows)
    if total_rows == 0:
        return

    table = doc.add_table(rows=total_rows, cols=n)
    table.style = "Table Grid"

    row_idx = 0

    if header_lines:
        hrow = table.rows[row_idx]; row_idx += 1
        for ci in range(n):
            cell = hrow.cells[ci]
            cell.paragraphs[0].clear()
            h_words = [w for ln in header_lines
                       for w in ln if _assign_col(w["x0"], col_starts) == ci]
            for w in h_words:
                fname = w.get("fontname", "")
                run = cell.paragraphs[0].add_run(
                    _remap_pua(w["text"], fname) + " ")
                run.bold = True
                run.font.name = _map_fontname(fname)
                run.font.size = Pt(float(w.get("size", 10)))

    for row_lines in logical_rows:
        drow = table.rows[row_idx]; row_idx += 1
        for ci in range(n):
            cell = drow.cells[ci]
            cell.paragraphs[0].clear()
            d_words = [w for ln in row_lines
                       for w in ln if _assign_col(w["x0"], col_starts) == ci]
            for w in d_words:
                fname = w.get("fontname", "")
                run = cell.paragraphs[0].add_run(
                    _remap_pua(w["text"], fname) + " ")
                run.font.name = _map_fontname(fname)
                run.font.size = Pt(float(w.get("size", 10)))


def _write_image(doc: Document, img: dict, page_width_pt: float) -> None:
    """Embed an image inline at its approximate width ratio."""
    result = _extract_image_bytes(img)
    if result is None:
        return
    raw, _ = result
    img_w_pt = float(img["x1"]) - float(img["x0"])
    width_in = min(img_w_pt / 72.0, 6.0)
    try:
        para = doc.add_paragraph()
        run  = para.add_run()
        run.add_picture(io.BytesIO(raw), width=Inches(width_in))
    except Exception:
        pass


def _write_header_layout(doc: Document, img: dict,
                          right_paras: list[list[list[dict]]], pw: float) -> None:
    """
    2-column borderless layout table: image in left cell, text paragraphs in right.
    Used when the PDF logo sits side-by-side with header text lines.
    """
    tbl = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(tbl)
    row = tbl.rows[0]

    # Left cell: image
    left = row.cells[0]
    _clear_cell_borders(left)
    result = _extract_image_bytes(img)
    if result:
        raw, _ = result
        img_w_pt = float(img["x1"]) - float(img["x0"])
        w_in = min(img_w_pt / 72.0, 2.0)
        left.paragraphs[0].add_run().add_picture(io.BytesIO(raw), width=Inches(w_in))

    # Right cell: text paragraphs
    right = row.cells[1]
    _clear_cell_borders(right)
    default_para = right.paragraphs[0]
    default_para._element.getparent().remove(default_para._element)
    for para_lines in right_paras:
        p = right.add_paragraph()
        _write_para_lines(p, para_lines)


# ---------------------------------------------------------------------------
# Per-page processing
# ---------------------------------------------------------------------------

def _process_page(doc: Document, page) -> None:
    pw = float(page.width)

    all_words = page.extract_words(extra_attrs=["size", "fontname"])

    # ── Detect content boxes ─────────────────────────────────────────────
    boxes = _content_boxes(page)

    # ── Assign words to boxes vs free ─────────────────────────────────────
    box_words: dict[int, list[dict]] = {i: [] for i in range(len(boxes))}
    free_words: list[dict] = []
    for w in all_words:
        wx1 = w.get("x1", w["x0"] + 1)
        wb  = w.get("bottom", w["top"] + 1)
        assigned = False
        for i, box in enumerate(boxes):
            if (box["x0"] - 2 <= w["x0"] and wx1 <= box["x1"] + 2 and
                    box["top"] - 2 <= w["top"] and wb <= box["bottom"] + 2):
                box_words[i].append(w)
                assigned = True
                break
        if not assigned:
            free_words.append(w)

    # ── Table detection (free words only) ────────────────────────────────
    sections = _section_rects(page, bordered_boxes=boxes)
    table_regions: list[tuple] = []
    for (y_top, y_bot) in sections:
        band = [w for w in free_words if y_top - 1 <= w["top"] <= y_bot + 1]
        if not band:
            continue
        lines = _group_into_lines(band)
        if not lines:
            continue
        cs = _detect_col_starts(lines[0])
        if len(cs) >= _MIN_TABLE_COLS:
            table_regions.append((y_top, y_bot, cs))

    non_table_free = [
        w for w in free_words
        if not any(yt - 1 <= w["top"] <= yb + 1 for (yt, yb, _cs) in table_regions)
    ]

    # ── Detect heading underlines ─────────────────────────────────────────
    ul_ys = _underline_ys(page)

    def _near_underline(top: float, bottom: float) -> bool:
        return any(top - 2 <= y <= bottom + 20 for y in ul_ys)

    # ── Box pairing ───────────────────────────────────────────────────────
    pairs, singles = _find_box_pairs(boxes)

    def _label_above(box_top: float, side: str = "left") -> str:
        candidates = [
            w for w in free_words
            if 0 < box_top - w["top"] < 35 and "symbol" not in w.get("fontname", "").lower()
        ]
        if not candidates:
            return ""
        ls = _group_into_lines(candidates)
        if not ls:
            return ""
        last_line = ls[-1]
        mid = pw / 2
        ws = [w for w in last_line if w["x0"] < mid] if side == "left" \
             else [w for w in last_line if w["x0"] >= mid]
        return " ".join(_remap_pua(w["text"], w.get("fontname", "")) for w in ws)

    # ── Label line exclusion ──────────────────────────────────────────────
    label_line_ys: set[float] = set()
    for (lbox, rbox) in pairs:
        for w in free_words:
            if 0 < lbox["top"] - w["top"] < 35:
                label_line_ys.add(round(w["top"], 1))
    for box in singles:
        for w in free_words:
            if 0 < box["top"] - w["top"] < 35:
                label_line_ys.add(round(w["top"], 1))

    # ── Group content lines into logical paragraphs ───────────────────────
    content_words = [
        w for w in non_table_free
        if round(w["top"], 1) not in label_line_ys
    ]
    content_lines = _group_into_lines(content_words)
    para_groups   = _group_into_paragraphs(content_lines, pw)

    # ── Images: detect side-by-side with text (header layout) ────────────
    images = list(page.images)
    consumed_para_ids: set[int] = set()

    image_events: list[tuple[float, str, object]] = []
    for img in images:
        img_top = float(img["top"])
        img_bot = float(img["bottom"])
        img_x1  = float(img["x1"])
        # Paragraphs that overlap this image's Y range AND start to the right of it
        side_paras = [
            pg for pg in para_groups
            if (pg[0][0]["top"] < img_bot + 5 and
                max(w.get("bottom", w["top"] + 12) for ln in pg for w in ln) > img_top - 5 and
                min(w["x0"] for ln in pg for w in ln) > img_x1 - 5)
        ]
        if side_paras:
            for pg in side_paras:
                consumed_para_ids.add(id(pg))
            image_events.append((img_top, "header", (img, side_paras)))
        else:
            image_events.append((img_top, "image", img))

    # ── Build full event list ─────────────────────────────────────────────
    rendered_table: set = set()
    rendered_box:   set = set()
    rendered_pair:  set = set()

    events: list[tuple[float, str, object]] = list(image_events)

    for (yt, yb, cs) in table_regions:
        events.append((yt, "table", (yt, yb, cs)))

    for pair_idx, (lbox, rbox) in enumerate(pairs):
        events.append((min(lbox["top"], rbox["top"]), "pair", pair_idx))

    for box_idx, box in enumerate(singles):
        events.append((box["top"], "single", box_idx))

    for pg in para_groups:
        events.append((pg[0][0]["top"], "para", pg))

    events.sort(key=lambda e: e[0])

    # ── Emit DOCX elements in reading order ───────────────────────────────
    bullet_buffer: list[list[dict]] = []

    def flush_bullets():
        for item in bullet_buffer:
            _write_bullet(doc, item)
        bullet_buffer.clear()

    for (y, kind, data) in events:

        if kind == "header":
            flush_bullets()
            img, side_paras = data
            _write_header_layout(doc, img, side_paras, pw)

        elif kind == "image":
            flush_bullets()
            _write_image(doc, data, pw)

        elif kind == "table":
            (yt, yb, cs) = data
            if (yt, yb) in rendered_table:
                continue
            rendered_table.add((yt, yb))
            flush_bullets()
            tbl_words = [w for w in free_words if yt - 1 <= w["top"] <= yb + 1]
            if tbl_words:
                _write_table(doc, tbl_words, cs, yt)

        elif kind == "pair":
            pair_idx = data
            if pair_idx in rendered_pair:
                continue
            rendered_pair.add(pair_idx)
            flush_bullets()
            lbox, rbox = pairs[pair_idx]
            ll = _label_above(lbox["top"], "left")
            rl = _label_above(rbox["top"], "right")
            _write_side_by_side(
                doc,
                box_words[boxes.index(lbox)], lbox,
                box_words[boxes.index(rbox)], rbox,
                ll, rl,
            )

        elif kind == "single":
            box_idx = data
            if box_idx in rendered_box:
                continue
            rendered_box.add(box_idx)
            flush_bullets()
            box = singles[box_idx]
            _write_code_box(doc, box_words[boxes.index(box)], box)

        elif kind == "para":
            pg: list[list[dict]] = data
            if id(pg) in consumed_para_ids:
                continue

            first_line = pg[0]

            # Bullet paragraph
            if first_line and _is_bullet_word(first_line[0]):
                for line in pg:
                    bullet_buffer.append(line)
                continue

            flush_bullets()

            all_bold      = all(_is_bold(w.get("fontname", "")) for w in first_line)
            top_y         = first_line[0]["top"]
            bot_y         = max(w.get("bottom", w["top"] + 12) for w in first_line)
            is_underlined = _near_underline(top_y, bot_y)

            if all_bold and is_underlined:
                # Section heading: bold Normal paragraph with bottom border
                para = doc.add_paragraph(style="Normal")
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after  = Pt(2)
                _write_runs(para, _merge_into_runs(first_line))
                _add_para_bottom_border(para)
                for extra in pg[1:]:
                    ep = doc.add_paragraph(style="Normal")
                    _write_para_lines(ep, [extra])
            else:
                # Normal paragraph: merge all wrapped lines
                para = doc.add_paragraph(style="Normal")
                _write_para_lines(para, pg)

    flush_bullets()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def convert(pdf_path: str | Path, title: str | None = None) -> bytes:
    """Convert a PDF to DOCX and return the raw bytes."""
    pdf_path  = Path(pdf_path)
    doc_title = title or pdf_path.stem

    doc = Document()

    # Document-level defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.core_properties.title = doc_title

    with pdfplumber.open(pdf_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            if page_idx > 0:
                para = doc.add_paragraph()
                run  = para.add_run()
                run.add_break(
                    __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
                )
            _process_page(doc, page)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys as _sys

    if len(_sys.argv) < 2:
        print("Usage: python pdf_to_docx.py <input.pdf> [output.docx]",
              file=_sys.stderr)
        _sys.exit(1)

    src = Path(_sys.argv[1])
    dst = Path(_sys.argv[2]) if len(_sys.argv) > 2 else src.with_suffix(".docx")
    dst.write_bytes(convert(src))
    print(f"Written → {dst}")
